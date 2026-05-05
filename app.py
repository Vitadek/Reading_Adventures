"""
Reading Adventures — a personal "parent hack" for generating
open-ended reading worksheets from EPUBs for a young reader.
"""
import json
import logging
import os
import re
from datetime import datetime
from pathlib import Path

import ebooklib
import psycopg2
from bs4 import BeautifulSoup
from ebooklib import epub
from flask import (
    Flask, abort, flash, redirect, render_template, request, url_for,
)
from openai import OpenAI, OpenAIError
from psycopg2.extras import Json, RealDictCursor
from werkzeug.utils import secure_filename

# --------------------------------------------------------------------------
# Configuration
# --------------------------------------------------------------------------
UPLOAD_FOLDER = Path("/app/uploads")
ALLOWED_EXTENSIONS = {"epub", "pdf", "docx"}
MAX_UPLOAD_BYTES = 50 * 1024 * 1024  # 50 MB — generous for any book format
MAX_BOOK_WORDS = int(os.environ.get("MAX_BOOK_WORDS", "100000"))
OPENAI_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
log = logging.getLogger("reading-adventures")

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = str(UPLOAD_FOLDER)
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_BYTES
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev-secret-change-me")

UPLOAD_FOLDER.mkdir(parents=True, exist_ok=True)


# --------------------------------------------------------------------------
# Database
# --------------------------------------------------------------------------
def get_db():
    """Open a new psycopg2 connection. Caller is responsible for closing."""
    return psycopg2.connect(
        host=os.environ["DB_HOST"],
        port=os.environ.get("DB_PORT", 5432),
        dbname=os.environ["DB_NAME"],
        user=os.environ["DB_USER"],
        password=os.environ["DB_PASSWORD"],
    )


# --------------------------------------------------------------------------
# Multi-format text extraction (EPUB, PDF, DOCX)
# --------------------------------------------------------------------------
def _normalize(text: str) -> str:
    """Collapse whitespace and strip."""
    return re.sub(r"\s+", " ", text or "").strip()


def _clean_html_to_text(html_bytes: bytes) -> str:
    """Strip HTML/CSS/scripts; return whitespace-normalized plain text."""
    soup = BeautifulSoup(html_bytes, "lxml")
    for tag in soup(["script", "style", "nav", "header", "footer"]):
        tag.decompose()
    return _normalize(soup.get_text(separator=" "))


def _read_epub(filepath: str):
    """Return (chunk_iterator, detected_title) for an EPUB."""
    book = epub.read_epub(filepath)
    title = None
    try:
        meta = book.get_metadata("DC", "title")
        if meta:
            title = (meta[0][0] or "").strip() or None
    except Exception:
        pass

    def chunks():
        for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
            text = _clean_html_to_text(item.get_content())
            if text:
                yield text

    return chunks(), title


def _read_pdf(filepath: str):
    """Return (chunk_iterator, detected_title) for a PDF (one chunk per page)."""
    from pypdf import PdfReader  # lazy import — only loaded if a PDF is uploaded
    reader = PdfReader(filepath)

    title = None
    try:
        if reader.metadata and reader.metadata.title:
            title = (reader.metadata.title or "").strip() or None
    except Exception:
        pass

    def chunks():
        for page_num, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception as e:
                log.warning("PDF page %d extraction failed: %s", page_num, e)
                continue
            text = _normalize(text)
            if text:
                yield text

    return chunks(), title


def _read_docx(filepath: str):
    """Return (chunk_iterator, detected_title) for a .docx file."""
    from docx import Document  # lazy import
    doc = Document(filepath)

    title = None
    try:
        if doc.core_properties.title:
            title = (doc.core_properties.title or "").strip() or None
    except Exception:
        pass

    def chunks():
        for para in doc.paragraphs:
            text = _normalize(para.text)
            if text:
                yield text
        # Walk tables too — some kids' books / activity books use them.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = _normalize(cell.text)
                    if text:
                        yield text

    return chunks(), title


_READERS = {
    ".epub": _read_epub,
    ".pdf":  _read_pdf,
    ".docx": _read_docx,
}


def extract_text_from_file(filepath: str, max_words: int = MAX_BOOK_WORDS):
    """
    Read a supported book file and return (full_text, detected_title).
    Stops accumulating chunks once max_words is reached.

    Supported formats: .epub, .pdf, .docx.
    Raises ValueError for unknown extensions.
    """
    ext = Path(filepath).suffix.lower()
    reader = _READERS.get(ext)
    if reader is None:
        raise ValueError(f"Unsupported file type: {ext}")

    chunk_iter, title = reader(filepath)

    collected: list[str] = []
    word_count = 0
    for text in chunk_iter:
        collected.append(text)
        word_count += len(text.split())
        if word_count >= max_words:
            break

    full_text = "\n\n".join(collected)
    # Hard cap (last chunk may have pushed us over).
    words = full_text.split()
    if len(words) > max_words:
        full_text = " ".join(words[:max_words])
    return full_text, title


# --------------------------------------------------------------------------
# OpenAI quiz generation
# --------------------------------------------------------------------------
SYSTEM_PROMPT = """You are a friendly, encouraging guide for a 7-year-old reader. Based on the provided book text, generate exactly 5 strictly OPEN-ENDED discussion questions, plus 1 fun drawing prompt at the end.

RULES:
1. NO yes/no questions.
2. NO multiple-choice questions.
3. NO simple fact-recall or trivia (e.g., do not ask 'What is the character's name?').
4. Focus entirely on empathy, imagination, and problem-solving. Ask things like 'How do you think the character felt when...', 'Why do you think they made that choice?', or 'What would you have done differently?'.
5. The 6th item must be a drawing prompt related to a scene in the book (e.g., 'Draw what you think the robot's island looks like!').

Use warm, simple language a 7-year-old can read on their own.

Return your output as a single JSON object with this exact shape:
{"questions": ["question 1", "question 2", "question 3", "question 4", "question 5", "drawing prompt"]}

Do not include any other keys, commentary, or markdown. The "questions" array MUST have exactly 6 strings, with the 6th being the drawing prompt."""


def generate_quiz_questions(book_text: str, max_attempts: int = 2) -> list[str]:
    """
    Call the OpenAI API and return a list of 6 strings (5 questions + 1 drawing prompt).
    Retries once if the response can't be parsed into the expected shape.
    """
    client = OpenAI()  # picks up OPENAI_API_KEY from env

    last_err: Exception | None = None
    for attempt in range(1, max_attempts + 1):
        try:
            log.info(
                "Calling OpenAI (model=%s, attempt=%d, words=%d)",
                OPENAI_MODEL, attempt, len(book_text.split()),
            )
            response = client.chat.completions.create(
                model=OPENAI_MODEL,
                response_format={"type": "json_object"},
                temperature=0.7,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {
                        "role": "user",
                        "content": f"Here is the book text:\n\n{book_text}",
                    },
                ],
            )
            content = response.choices[0].message.content or ""
            data = json.loads(content)
            questions = data.get("questions")

            if (
                isinstance(questions, list)
                and len(questions) == 6
                and all(isinstance(q, str) and q.strip() for q in questions)
            ):
                return [q.strip() for q in questions]

            last_err = ValueError(
                f"Malformed shape: expected 6 non-empty strings, got {questions!r}"
            )
            log.warning("Attempt %d returned malformed JSON: %s", attempt, last_err)
        except (json.JSONDecodeError, OpenAIError, KeyError) as e:
            last_err = e
            log.warning("Attempt %d failed: %s", attempt, e)

    raise RuntimeError(f"Failed to generate quiz after {max_attempts} attempts: {last_err}")


# --------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------
def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


# --------------------------------------------------------------------------
# Routes
# --------------------------------------------------------------------------
@app.route("/", methods=["GET"])
def index():
    with get_db() as conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute(
            """
            SELECT b.id,
                   b.title,
                   b.filename,
                   b.uploaded_at,
                   COUNT(q.id) AS quiz_count,
                   MAX(q.id)   AS latest_quiz_id
            FROM books b
            LEFT JOIN quizzes q ON q.book_id = b.id
            GROUP BY b.id
            ORDER BY b.uploaded_at DESC
            """
        )
        books = cur.fetchall()
    return render_template("index.html", books=books)


@app.route("/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        flash("No file part in the form.", "error")
        return redirect(url_for("index"))

    file = request.files["file"]
    if not file or not file.filename:
        flash("Please choose a file to upload.", "error")
        return redirect(url_for("index"))

    if not allowed_file(file.filename):
        flash("Only .epub, .pdf, and .docx files are supported.", "error")
        return redirect(url_for("index"))

    safe_name = secure_filename(file.filename)
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    stored_name = f"{timestamp}_{safe_name}"
    filepath = UPLOAD_FOLDER / stored_name
    file.save(filepath)

    # Try reading the title from file metadata; fall back to filename.
    title_fallback = safe_name.rsplit(".", 1)[0].replace("_", " ")
    try:
        _, detected_title = extract_text_from_file(str(filepath), max_words=200)
        title = detected_title or title_fallback
    except Exception as e:
        log.warning("Could not read metadata from %s: %s", stored_name, e)
        title = title_fallback

    with get_db() as conn, conn.cursor() as cur:
        cur.execute(
            "INSERT INTO books (title, filename) VALUES (%s, %s)",
            (title[:500], stored_name),
        )
        conn.commit()

    flash(f"Uploaded \u201C{title}\u201D successfully.", "success")
    return redirect(url_for("index"))


@app.route("/generate/<int:book_id>", methods=["POST"])
def generate(book_id: int):
    with get_db() as conn:
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute("SELECT * FROM books WHERE id = %s", (book_id,))
            book = cur.fetchone()
        if not book:
            abort(404)

        filepath = UPLOAD_FOLDER / book["filename"]
        if not filepath.exists():
            flash("That book file is missing on disk.", "error")
            return redirect(url_for("index"))

        try:
            text, _ = extract_text_from_file(str(filepath))
            if not text.strip():
                flash(
                    "Couldn't extract any readable text from this file. "
                    "If it's a scanned PDF, OCR isn't supported yet.",
                    "error",
                )
                return redirect(url_for("index"))
            questions = generate_quiz_questions(text)
        except Exception as e:
            log.exception("Quiz generation failed for book %s", book_id)
            flash(f"Quiz generation failed: {e}", "error")
            return redirect(url_for("index"))

        with conn.cursor() as cur:
            cur.execute(
                "INSERT INTO quizzes (book_id, questions_json) VALUES (%s, %s) RETURNING id",
                (book_id, Json(questions)),
            )
            quiz_id = cur.fetchone()[0]
            conn.commit()

    return redirect(url_for("view_quiz", quiz_id=quiz_id))


@app.route("/quiz/<int:quiz_id>", methods=["GET"])
def view_quiz(quiz_id: int):
    with get_db() as conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute(
            """
            SELECT q.id            AS quiz_id,
                   q.questions_json AS questions,
                   q.created_at,
                   b.id            AS book_id,
                   b.title         AS book_title
            FROM quizzes q
            JOIN books b ON b.id = q.book_id
            WHERE q.id = %s
            """,
            (quiz_id,),
        )
        quiz = cur.fetchone()
    if not quiz:
        abort(404)
    return render_template("quiz.html", quiz=quiz)


@app.route("/delete/<int:book_id>", methods=["POST"])
def delete_book(book_id: int):
    """Delete a book, its EPUB file, and all its quizzes (via FK CASCADE)."""
    with get_db() as conn, conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute("SELECT title, filename FROM books WHERE id = %s", (book_id,))
        book = cur.fetchone()
        if not book:
            abort(404)
        # ON DELETE CASCADE on quizzes.book_id removes related quizzes for free.
        cur.execute("DELETE FROM books WHERE id = %s", (book_id,))
        conn.commit()

    # Best-effort file removal — don't fail the request if it's already gone.
    filepath = UPLOAD_FOLDER / book["filename"]
    try:
        filepath.unlink(missing_ok=True)
    except OSError as e:
        log.warning("Could not delete file %s: %s", filepath, e)

    flash(f"Removed \u201C{book['title']}\u201D and any related worksheets.", "success")
    return redirect(url_for("index"))


@app.route("/healthz")
def healthz():
    """Lightweight health endpoint for docker/uptime checks."""
    try:
        with get_db() as conn, conn.cursor() as cur:
            cur.execute("SELECT 1")
            cur.fetchone()
        return {"status": "ok"}, 200
    except Exception as e:
        return {"status": "error", "detail": str(e)}, 500


# --------------------------------------------------------------------------
# Error handlers
# --------------------------------------------------------------------------
@app.errorhandler(413)
def too_large(_e):
    flash(f"That file is too big. Max size is {MAX_UPLOAD_BYTES // (1024*1024)} MB.", "error")
    return redirect(url_for("index"))


@app.errorhandler(404)
def not_found(_e):
    return render_template("base.html", not_found=True), 404


if __name__ == "__main__":
    # For local non-docker development only.
    app.run(host="0.0.0.0", port=5000, debug=True)
